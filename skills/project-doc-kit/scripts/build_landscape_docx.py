from __future__ import annotations

import argparse
import re
import os
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


DEFAULT_ORDER = [
    "README.md",
    "01-onboarding-and-context.md",
    "02-architecture-decomposition.md",
    "03-services-and-modules.md",
    "04-data-model-and-storage.md",
    "05-business-flows-sequence.md",
    "06-api-and-integration-contracts.md",
    "07-frontend-backend-code-map.md",
    "08-devops-testing-observability.md",
    "09-git-history-and-delivery-map.md",
    "10-problems-and-contradictions.md",
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a landscape DOCX project documentation album.")
    parser.add_argument("docs_dir", help="Target docs/YYYYMMDD directory.")
    parser.add_argument("--date", help="Date suffix. Defaults to docs directory name.")
    return parser.parse_args()


def ordered_markdown_files(docs_dir: Path) -> list[Path]:
    ordered = [docs_dir / name for name in DEFAULT_ORDER if (docs_dir / name).exists()]
    known = {path.name for path in ordered}
    extra = sorted(
        path
        for path in docs_dir.glob("*.md")
        if path.name not in known and not path.name.startswith("project-documentation-")
    )
    return ordered + extra


def build_combined_markdown(docs_dir: Path, date_suffix: str) -> str:
    parts = [
        f"# Project Documentation Kit {date_suffix}",
        "",
        "Landscape DOCX version. Includes current implementation debt and contradictions.",
        "",
    ]

    for index, path in enumerate(ordered_markdown_files(docs_dir)):
        text = path.read_text(encoding="utf-8").strip()
        if index > 0:
            parts.extend(["\\pagebreak", ""])
        parts.extend([text, ""])

    return "\n".join(parts).strip() + "\n"


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)

    for level in range(1, 5):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)

    doc.styles["Heading 1"].font.size = Pt(18)
    doc.styles["Heading 2"].font.size = Pt(15)
    doc.styles["Heading 3"].font.size = Pt(12)
    doc.styles["Heading 4"].font.size = Pt(11)
    return doc


def find_mmdc() -> str | None:
    npm_global = Path(os.environ.get("APPDATA", "")) / "npm"
    for name in ("mmdc.cmd", "mmdc.ps1", "mmdc"):
        candidate = npm_global / name
        if candidate.exists():
            return str(candidate)
    return shutil.which("mmdc.cmd") or shutil.which("mmdc")


def render_mermaid_to_png(mermaid_code: str, diagram_index: int, cache_dir: Path) -> Path | None:
    mmdc = find_mmdc()
    if not mmdc:
        print(f"warning: mmdc not found; diagram {diagram_index} will be inserted as code")
        return None

    cache_dir.mkdir(parents=True, exist_ok=True)
    source = cache_dir / f"diagram_{diagram_index}.mmd"
    output = cache_dir / f"diagram_{diagram_index}.png"
    source.write_text(mermaid_code, encoding="utf-8")

    try:
        result = subprocess.run(
            [mmdc, "-i", str(source), "-o", str(output), "-b", "white", "-w", "1800", "-s", "2"],
            capture_output=True,
            text=True,
            timeout=90,
            shell=True,
        )
    except subprocess.TimeoutExpired:
        print(f"warning: mmdc timed out for diagram {diagram_index}")
        return None
    except Exception as exc:
        print(f"warning: mmdc failed for diagram {diagram_index}: {exc}")
        return None

    if output.exists() and output.stat().st_size > 0:
        print(f"rendered diagram {diagram_index}: {output.stat().st_size} bytes")
        return output

    stderr = (result.stderr or "").strip()
    if stderr:
        print(f"warning: mmdc produced no image for diagram {diagram_index}: {stderr[:240]}")
    else:
        print(f"warning: mmdc produced no image for diagram {diagram_index}")
    return None


def clean_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.replace("**", "").replace("`", "")


def add_paragraph(doc: Document, text: str, style: str = "Normal"):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.add_run(clean_inline(text))
    return paragraph


def parse_table_line(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator(row: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in row)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for col_index in range(cols):
            value = row[col_index] if col_index < len(row) else ""
            cell = cells[col_index]
            cell.text = clean_inline(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9)
                    if row_index == 0:
                        run.font.bold = True
    doc.add_paragraph()


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    if not code_lines:
        return
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.5)
    run = paragraph.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8)


def add_mermaid_image(doc: Document, mermaid_lines: list[str], diagram_index: int, cache_dir: Path) -> None:
    mermaid_code = "\n".join(mermaid_lines)
    png = render_mermaid_to_png(mermaid_code, diagram_index, cache_dir)
    if not png:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("[ Diagram failed to render, see the source Mermaid code below ]")
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = RGBColor(150, 80, 80)
        add_code_block(doc, mermaid_lines)
        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    try:
        from PIL import Image

        with Image.open(png) as image:
            width_px, height_px = image.size
        aspect = height_px / width_px
        width_cm = 25.5
        height_cm = width_cm * aspect
        if height_cm > 15.5:
            height_cm = 15.5
            width_cm = height_cm / aspect
        paragraph.add_run().add_picture(str(png), width=Cm(width_cm))
    except Exception:
        paragraph.add_run().add_picture(str(png), width=Cm(24))

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(f"Figure {diagram_index}")
    run.font.size = Pt(9)
    run.italic = True


def convert_markdown_to_docx(markdown: str, output: Path) -> Path:
    doc = setup_document()
    lines = markdown.splitlines()
    in_code = False
    code_language = ""
    code_lines: list[str] = []
    table_rows: list[list[str]] = []
    diagram_index = 0
    mermaid_cache_dir = output.parent / ".docx_mermaid_cache"

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            rows = [row for row in table_rows if not is_separator(row)]
            add_table(doc, rows)
            table_rows = []

    def flush_code() -> None:
        nonlocal code_lines
        add_code_block(doc, code_lines)
        code_lines = []

    for raw_line in lines:
        line = raw_line.rstrip()

        if line.startswith("```"):
            if in_code:
                if code_language == "mermaid":
                    diagram_index += 1
                    add_mermaid_image(doc, code_lines, diagram_index, mermaid_cache_dir)
                    code_lines = []
                else:
                    flush_code()
                in_code = False
                code_language = ""
            else:
                flush_table()
                in_code = True
                code_language = line[3:].strip().lower()
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.strip() == "\\pagebreak":
            flush_table()
            doc.add_page_break()
            continue

        if line.startswith("|") and line.endswith("|"):
            table_rows.append(parse_table_line(line))
            continue

        flush_table()

        if not line.strip():
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            add_paragraph(doc, heading.group(2), f"Heading {len(heading.group(1))}")
            continue

        bullet = re.match(r"^(\s*)[-*]\s+(.+)$", line)
        if bullet:
            paragraph = add_paragraph(doc, "• " + bullet.group(2))
            paragraph.paragraph_format.left_indent = Cm(0.6 + len(bullet.group(1)) * 0.1)
            continue

        if line.startswith(">"):
            paragraph = add_paragraph(doc, line.lstrip("> "))
            paragraph.paragraph_format.left_indent = Cm(0.8)
            if paragraph.runs:
                paragraph.runs[0].italic = True
            continue

        paragraph = add_paragraph(doc, line)
        if len(line) < 90 and line.endswith(":") and paragraph.runs:
            paragraph.runs[0].bold = True

    flush_table()
    flush_code()

    if doc.paragraphs:
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        doc.save(output)
        saved_path = output
    except PermissionError:
        saved_path = output.with_name(f"{output.stem}-new{output.suffix}")
        print(f"warning: {output.name} is locked; saving as {saved_path.name}")
        doc.save(saved_path)
    shutil.rmtree(mermaid_cache_dir, ignore_errors=True)
    return saved_path


def main() -> None:
    args = parse_args()
    docs_dir = Path(args.docs_dir).resolve()
    if not docs_dir.exists():
        raise SystemExit(f"docs_dir does not exist: {docs_dir}")

    date_suffix = args.date or docs_dir.name
    combined = docs_dir / f"project-documentation-{date_suffix}-full.md"
    docx = docs_dir / f"project-documentation-{date_suffix}-landscape.docx"

    markdown = build_combined_markdown(docs_dir, date_suffix)
    combined.write_text(markdown, encoding="utf-8")
    saved_docx = convert_markdown_to_docx(markdown, docx)

    print(f"markdown={combined}")
    print(f"docx={saved_docx}")


if __name__ == "__main__":
    main()
